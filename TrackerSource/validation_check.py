import streamlit as st
import json
from datetime import date
from docx import Document

# Predefined standards and test cases
standards_map = {
    "Railway - Inverter": ["IEC 62040-3", "IEC 61375", "IEC 62236", "EN 50155", "EN 50121"],
    "Railway - Converter": ["IEC 62040-3", "IEC 61000-4-17", "EN 50121-3-2", "EN 50121"],
    "Industrial - Battery Charger": ["IEC 62040-3", "IEC 61010-1", "IEC 62040-3"],
    "Industrial - Inverter": ["IEC 62040-3", "IEC 61800-3", "IEC 61000-6-2"],
    "Industrial - Converter": ["IEC 62040-3", "IEC 61000-4-5", "IEC 61000-3-2"]
}

predefined_tests = [
    {"id": "TC001 - INPUT VOLTAGE SWEEP", "objective": "Sweep Vin * 0.6 & Vin * 1.4, plus sweep in operational range"},
    {"id": "TC002 - VOLTAGE INTERRUPT TEST", "objective": "Interrupt Vin for 10 ms, 100 ms, 500 ms and check if Vout recovers without resets"},
    {"id": "TC003 - OVERVOLTAGE PROTECTION", "objective": "Vin > Vin,max - units protect or shut down"},
    {"id": "TC004 - UNDERVOLTAGE PROTECTION", "objective": "Vin < Vin,min - units protect or shut down"},
    {"id": "TC005 - THERMAL STRESS", "objective": "Performance check for 1 hr, check the temperatures"},
    {"id": "TC006 - LOAD TRANSIENT RESPONSE", "objective": "Evaluate Vout response to load changes"},
    {"id": "TC007 - EMC TEST", "objective": "ORing position, the MOSFET is always ON, only radiated done"}
]

class ValidationChecker:
    def __init__(self):
        self.metadata = {}
        self.test_cases = []

    def parse_docx(self, file):
        doc = Document(file)
        metadata = {}
        test_cases = []

        for table in doc.tables:
            if len(table.columns) == 2:
                for row in table.rows:
                    key = row.cells[0].text.strip().lower().replace(" ", "_")
                    value = row.cells[1].text.strip()
                    metadata[key] = value
                break

        for table in doc.tables:
            if len(table.columns) == 3 and table.rows[0].cells[0].text.strip() == "Test Case ID":
                for row in table.rows[1:]:
                    test_cases.append({
                        "id": row.cells[0].text.strip(),
                        "objective": row.cells[1].text.strip(),
                        "result": row.cells[2].text.strip()
                    })
                break
        return metadata, test_cases

    def run(self):
        st.title("Validation Plan Generator")

        uploaded_doc = st.file_uploader("📤 Upload a .docx file to populate the form", type=["docx"])
        if uploaded_doc:
            self.metadata, self.test_cases = self.parse_docx(uploaded_doc)
            st.success("Document parsed successfully.")

        col1, col2 = st.columns([2, 1])
        with col1:
            with st.form("validation_form"):
                st.subheader("Component Information")
                c1, c2 = st.columns(2)
                self.metadata["project_part"] = c1.text_input("Project Part Number", value=self.metadata.get("project_part", ""))
                self.metadata["component_change"] = c2.text_input("Component Change", value=self.metadata.get("component_change", ""))

                c3, c4 = st.columns(2)
                self.metadata["device_model"] = c3.text_input("Device Model", value=self.metadata.get("device_model", ""))
                self.metadata["input"] = c4.text_input("Input", value=self.metadata.get("input", ""))

                c5, c6 = st.columns(2)
                self.metadata["output"] = c5.text_input("Output", value=self.metadata.get("output", ""))
                self.metadata["efficiency"] = c6.text_input("Efficiency", value=self.metadata.get("efficiency", ""))

                self.metadata["product_type"] = st.selectbox("Product Type", list(standards_map.keys()), index=0)
                self.metadata["standards"] = ", ".join(standards_map.get(self.metadata["product_type"], []))

                c7, c8 = st.columns(2)
                self.metadata["environment"] = c7.text_input("Test Environment", value=self.metadata.get("environment", ""))
                self.metadata["engineer"] = c8.text_input("Engineer", value=self.metadata.get("engineer", ""))

                self.metadata["test_date"] = st.date_input("Test Date", value=date.today()).strftime("%Y-%m-%d")
                self.metadata["data_insertion"] = st.text_input("Data Insertion", value=self.metadata.get("data_insertion", ""))

                st.subheader("Test Cases")
                self.test_cases = []
                for i, test in enumerate(predefined_tests):
                    t1, t2 = st.columns([1, 3])
                    test_id = t1.text_input("Test ID", value=test["id"], key=f"id_{i}")
                    objective = t2.text_input("Objective", value=test["objective"], key=f"obj_{i}")
                    result = st.text_input("Result", value="", key=f"res_{i}")
                    self.test_cases.append({"id": test_id, "objective": objective, "result": result})

                submitted = st.form_submit_button("Generate Validation Plan")

        with col2:
            st.subheader("Live Preview")
            if self.metadata:
                st.text("Metadata:")
                st.json(self.metadata)
            if self.test_cases:
                st.text("Test Cases:")
                st.json(self.test_cases)

        if submitted:
            st.success("Validation plan generated successfully.")
            st.download_button("Download JSON", data=json.dumps({"metadata": self.metadata, "test_cases": self.test_cases}, indent=2),
                               file_name="validation_plan.json", mime="application/json")