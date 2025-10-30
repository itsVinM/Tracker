import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io, os
from datetime import date
from io import BytesIO

# Predefined comparison fields per component type
PRODUCT_COMPARISON_FIELDS = {
    "MOSFET": {
        "materiales": ["Vds", "Id", "Rds(on)", "Qg", "Vgs(th)", "Tj(max)", "Pd(max)", "Compliance"],
        "dimensionado": ["Package", "Pin Count", "Mounting Type", "Size"]
    },
    "Diode": {
        "materiales": ["Vr", "If", "Vf", "trr", "Tj(max)", "Pd(max)", "Compliance"],
        "dimensionado": ["Package", "Pin Count", "Mounting Type", "Size"]
    },
    "Inductor": {
        "materiales": ["Inductance", "Rated Current", "Saturation Current", "DCR", "Shielding", "Compliance"],
        "dimensionado": ["Core Size", "Height", "Footprint", "Mounting Type"]
    },
    "Connector": {
        "materiales": ["Current Rating", "Voltage Rating", "Contact Resistance", "Insulation Resistance", "Compliance"],
        "dimensionado": ["Pitch", "Rows", "Contact Count", "Mounting Type"]
    },
    "DC-DC Converter": {
        "materiales": ["Input Voltage Range", "Output Voltage", "Output Current", "Efficiency", "Isolation Voltage", "Compliance"],
        "dimensionado": ["Module Size", "Height", "Pin Count", "Mounting Type"]
    },
    "Capacitor": {
        "materiales": ["Capacitance", "Voltage Rating", "ESR", "Ripple Current", "Compliance"],
        "dimensionado": ["Package", "Size", "Mounting Type"]
    },
    "Custom": {
        "materiales": [],
        "dimensionado": []
    }
}

class HomologationApp:
    def __init__(self):
        if 'report_data' not in st.session_state:
            st.session_state.report_data = {}
    

    def add_hyperlink(self, paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)


    def display_form(self):
        logo_path = os.path.join(os.getcwd(), "premium_psu_logo.png")
        col1,col2=st.columns(2)
        with col1:
            data = st.session_state.report_data

            # General Information
            with st.expander("General Information", expanded=True):
                cols = st.columns(3)
                data['product_type'] = cols[0].selectbox("Component Type", list(PRODUCT_COMPARISON_FIELDS.keys()))
                data['doc_id'] = cols[1].text_input("Document ID", "H-2025-133")
                data['edition'] = cols[2].text_input("Edition", "2")

                cols2 = st.columns(3)
                data['codigos'] = cols2[0].text_input("Códigos", "26010206")
                data['date'] = cols2[1].date_input("Date", value=date.today()).strftime("%d.%m.%Y")
                data['author'] = cols2[2].text_input("Author", "V.Mocanu")

            # Objecto
            with st.expander("Objecto", expanded=True):
                data['objeto'] = st.text_area("Objecto", "Se estudia la posibilidad de homologar el componente...")
            
            with st.expander("Motivo", expanded=True):
                data['motivo'] = st.text_area("Motivo", 
                    "Solicitante: \n"
                    "Motivo: ")

            with st.expander("Investigativo", expanded=True):
                # Motivation text area
                data['investigativo'] = st.text_area(
                    "Investigativo",
                    "El componente que se compraba hasta ahora es el \n"
                    "Los diseños de los componentes se pueden consultar en la siguiente carpeta:\n"
                    "G:\\Laboratori\\PLANOS - M. PRIMAS_Backup\\Data sheets")

                # Component list with name + link together
                num_links = st.slider("Número de componentes a comparar", 1, 5, 2)
                data['datasheet_links'] = []

                
                st.markdown("### Componentes y enlaces")
                for i in range(num_links):
                    st.markdown(f"**Componente {i+1}**")
                    component_name = st.text_input(f"Nombre del componente {i+1}", key=f"ds_name_{i}", placeholder="Ej: ZUGO D12P154FST0001T")
                    component_link = st.text_input(f"Enlace del componente {i+1}", key=f"ds_url_{i}", placeholder="Ej: http://example.com/datasheet")

                    if not component_name.strip():
                        component_name = f"Component_{i+1}"

                    data['datasheet_links'].append({'name': component_name, 'url': component_link})

            # Comparison Tables
            with st.expander("Comparison Tables", expanded=True):
                # Electrical Properties
                st.markdown("#### Electrical Properties")
                comp_names = [comp['name'] if comp['name'].strip() else f"Component_{i+1}" for i, comp in enumerate(data['datasheet_links'])]
                materiales_df = pd.DataFrame(columns=["Field"] + comp_names)
                for field in PRODUCT_COMPARISON_FIELDS[data['product_type']]['materiales']:
                    row = {"Field": field}
                    for name in comp_names:
                        row[name] = ""
                    materiales_df.loc[len(materiales_df)] = row
                st.data_editor(materiales_df, num_rows="dynamic", key="materiales_editor")
                data['materiales'] = materiales_df.to_dict(orient="records")

                # Physical Dimensions
                st.markdown("#### Physical Dimensions")
                dimensionado_df = pd.DataFrame(columns=["Field"] + comp_names)
                for field in PRODUCT_COMPARISON_FIELDS[data['product_type']]['dimensionado']:
                    row = {"Field": field}
                    for name in comp_names:
                        row[name] = ""
                    dimensionado_df.loc[len(dimensionado_df)] = row
                st.data_editor(dimensionado_df, num_rows="dynamic", key="dimensionado_editor")
                data['dimensionado'] = dimensionado_df.to_dict(orient="records")

            # Conclusion
            with st.expander("Conclusion", expanded=True):
                data['conclusion'] = st.text_area("Conclusión", "El componente propuesto tiene un diseño con mismas dimensiones de las opciones homologadas.")
        with col2:
            st.subheader("📄 Live Preview")
            self.display_preview()
            
            if st.button("Generate DOCX Report"):
                self.generate_doc(data, logo_path)

    def generate_doc(self,data, logo_path):
        # Generate DOCX
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Aptos Narrow'
        font.size = Pt(11)

        # Header table
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Left: Logo
        cell_logo = table.cell(0, 0)
        try:
            run = cell_logo.paragraphs[0].add_run()
            run.add_picture(logo_path, width=Inches(1.2))
        except:
            cell_logo.text = "Logo not found"

        # Center
        cell_center = table.cell(0, 1)
        p_center = cell_center.paragraphs[0]
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_center = p_center.add_run(f"{data['product_type']}\nSolicitud de homologación\nCódigos: {data['codigos']}")
        run_center.bold = True

        # Right
        cell_right = table.cell(0, 2)
        info = (
            f"Doc ID: {data['doc_id']}\n"
            f"Date: {data['date']}\n"
            f"Author: {data['author']}\n"
            f"Edition: {data['edition']}\n"
            f"Comment: {data.get('comments', {}).get('Códigos', 'OK')}"
        )
        cell_right.text = info

        # Sections
        doc.add_paragraph()
        doc.add_heading('1. Objetivo', level=1)
        doc.add_paragraph(data['objeto'])
        doc.add_heading('2. Motivo de la solicitud', level=1)
        doc.add_paragraph(data.get('motivo', 'Texto explicativo aquí...'))
        doc.add_heading('3. Investigativo previo', level=1)
        doc.add_paragraph(data.get('investigativo', 'Detalles previos...'))
        doc.add_heading('4. Comparativa parámetros', level=1)
        doc.add_heading('Materiales y características mecánicas', level=2)

        # Materiales table
        materiales = data.get('materiales', [])
        if materiales:
            keys = list(materiales[0].keys())
            table_mat = doc.add_table(rows=1, cols=len(keys))
            table_mat.style = 'Table Grid'
            for i, key in enumerate(keys):
                cell = table_mat.cell(0, i)
                cell.text = key
                cell.paragraphs[0].runs[0].bold = True
            for row_data in materiales:
                row = table_mat.add_row().cells
                for i, key in enumerate(keys):
                    row[i].text = str(row_data.get(key, ''))

        doc.add_heading('Dimensiones', level=2)
        dimensionado = data.get('dimensionado', [])
        if dimensionado:
            keys = list(dimensionado[0].keys())
            table_dim = doc.add_table(rows=1, cols=len(keys))
            table_dim.style = 'Table Grid'
            for i, key in enumerate(keys):
                cell = table_dim.cell(0, i)
                cell.text = key
                cell.paragraphs[0].runs[0].bold = True
            for row_data in dimensionado:
                row = table_dim.add_row().cells
                for i, key in enumerate(keys):
                    row[i].text = str(row_data.get(key, ''))

        doc.add_heading('6. Conclusiones', level=1)
        doc.add_paragraph(data["conclusion"])

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # ✅ Add Download Button Here
        st.download_button(
            label="Download DOCX",
            data=buffer,
            file_name="Homologation_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
            

    def display_preview(self, logo_path=None):
        data = st.session_state.report_data

        # CSS Styles
        st.markdown(
            "<style>"
            ".doc-container { background-color: #fff; color: #000; font-family: Arial, sans-serif; padding: 20px; max-width: 900px; margin: auto; border: 1px solid #ccc; }"
            ".header-table { width: 100%; border-collapse: collapse; margin-bottom: 20px; }"
            ".header-table td { vertical-align: top; padding: 10px; }"
            ".header-left img { max-height: 60px; }"
            ".header-center { text-align: center; }"
            ".header-center h2 { margin: 5px 0; font-size: 24px; color: #000; }"
            ".header-center p { margin: 2px 0; font-size: 16px; }"
            ".header-right table { border-collapse: collapse; font-size: 14px; }"
            ".header-right td { padding: 4px 8px; }"
            ".ok-cell { background-color: #92D050; color: #000; font-weight: bold; }"
            ".nok-cell { background-color: #FF0000; color: #fff; font-weight: bold; }"
            ".section-title { font-size: 18px; font-weight: bold; color: #0070C0; margin-top: 20px; }"
            ".sub-section-title { font-size: 16px; font-weight: bold; color: #0070C0; margin-top: 15px; }"
            ".comparison-table { width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 14px; }"
            ".comparison-table th, .comparison-table td { border: 1px solid #000; padding: 6px; text-align: center; }"
            "</style>",
            unsafe_allow_html=True
        )

        # Header HTML
        header_html = (
            '<table class="header-table">'
            '<tr>'
            '<td class="header-left">'
            f'{f"<img src=\'{logo_path}\' alt=\'Logo\'>" if logo_path else ""}'
            '</td>'
            '<td class="header-center">'
            f'<h2>{data["product_type"]}</h2>'
            '<p><strong>Solicitud de homologación</strong></p>'
            f'<p>Códigos: {data["codigos"]}</p>'
            '</td>'
            '<td class="header-right">'
            '<table>'
            f'<tr><td>Doc ID:</td><td>{data["doc_id"]}</td></tr>'
            f'<tr><td>Edition:</td><td>{data["edition"]}</td></tr>'
            f'<tr><td>Date:</td><td>{data["date"]}</td></tr>'
            f'<tr><td>Author:</td><td>{data["author"]}</td></tr>'
            '</table>'
            '</td>'
            '</tr>'
            '</table>'
        )
        st.markdown(header_html, unsafe_allow_html=True)

        # Motivo section with proper spacing
        st.markdown('<div class="section-title">1. Objecto</div>', unsafe_allow_html=True)
        st.markdown(f"<div>{data['objeto'].replace('\n', '<br>')}</div>", unsafe_allow_html=True)

        # Motivo section with proper spacing
        st.markdown('<div class="section-title">2. Motivo de la solicitud</div>', unsafe_allow_html=True)
        st.markdown(f"<div>{data['motivo'].replace('\n', '<br>')}</div>", unsafe_allow_html=True)

        
        component_list = data.get('componentes', [])

        # Add component list if available
        if component_list:
            st.markdown("<strong>Componentes:</strong>", unsafe_allow_html=True)
            st.markdown("<ul>" + "".join([f"<li>{c}</li>" for c in component_list]) + "</ul>", unsafe_allow_html=True)

        # Display component list if available
        if component_list:
            st.markdown('<strong>Componentes:</strong>', unsafe_allow_html=True)
            st.write(", ".join(component_list))  # Or use bullet points:
            # st.markdown("<ul>" + "".join([f"<li>{c}</li>" for c in component_list]) + "</ul>", unsafe_allow_html=True)

        st.markdown('<div class="section-title">3. Investigativo previo</div>', unsafe_allow_html=True)
        formatted_investigativo = (data.get('investigativo', 'Detalles previos...')).replace('\n', '<br>')
        st.markdown(f"<div>{formatted_investigativo}</div>", unsafe_allow_html=True)

    
        
        datasheet_links = data.get('datasheet_links', [])
        components_html = ""
        if datasheet_links:
            components_html += "<br><strong>Componentes:</strong><ul>"
            for comp in datasheet_links:
                name = comp.get('name', 'Componente')
                url = comp.get('url', '')
                if url.strip():
                    # ✅ Embed link in the name
                    components_html += f'<li><{url}{name}</a></li>'
                else:
                    components_html += f'<li>{name}</li>'
            components_html += "</ul>"

        
        # Combine investigativo text and components list
        st.markdown(f"<div>{components_html}</div>", unsafe_allow_html=True)




        st.markdown('<div class="section-title">4. Comparativa parámetros</div>', unsafe_allow_html=True)

        # Subsection: Materiales y características mecánicas
        st.markdown('<div class="sub-section-title">Materiales y características mecánicas</div>', unsafe_allow_html=True)
        materiales_df = pd.DataFrame(data['materiales'])
        st.markdown(materiales_df.to_html(classes='comparison-table', index=False), unsafe_allow_html=True)

        # Subsection: Dimensiones
        st.markdown('<div class="sub-section-title">Dimensiones</div>', unsafe_allow_html=True)
        dimensionado_df = pd.DataFrame(data['dimensionado'])
        st.markdown(dimensionado_df.to_html(classes='comparison-table', index=False), unsafe_allow_html=True)

        # Conclusion
        st.markdown('<div class="section-title">6. Conclusiones</div>', unsafe_allow_html=True)
        st.write(data["conclusion"])

            
