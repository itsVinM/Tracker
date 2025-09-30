import os, re, shutil, json, sqlite3, plotly

import pandas as pd
from docx import Document
from datetime import datetime
import plotly.figure_factory as ff
import plotly.graph_objects as go
from sqlalchemy import create_engine
import matplotlib.pyplot as plt
import streamlit as st

def create_dut_chart(data, title, duration_days=3):
    # Rename columns for Gantt chart compatibility
    gantt_data = data.rename(columns={
        'START DATE': 'Start',
        'Used': 'Task'
    })

    # Ensure all entries in the Task column are strings
    gantt_data['Task'] = gantt_data['Task'].astype(str)

    # Convert Start to datetime and compute Finish using fixed duration
    gantt_data['Start'] = pd.to_datetime(gantt_data['Start'])
    gantt_data['Finish'] = gantt_data['Start'] + pd.to_timedelta(duration_days, unit='d')

    # Check if gantt_data is empty
    if gantt_data.empty:
        raise ValueError("No data available for the Gantt chart.")

    # Generate colors based on the number of unique tasks
    unique_tasks = gantt_data['Task'].nunique()
    cmap = plt.get_cmap('tab20')
    colors = [cmap(i) for i in range(unique_tasks)]
    colors = ['#%02x%02x%02x' % (int(r*255), int(g*255), int(b*255)) for r, g, b, _ in colors]

    # Create Gantt chart
    fig = ff.create_gantt(gantt_data, index_col='Task', show_colorbar=True, group_tasks=True, colors=colors[:unique_tasks])
    fig.update_layout(title_text=title, xaxis_title='Timeline', autosize=True)

    # Return chart as JSON
    return json.dumps(fig, cls=plotly.utils.PlotlyJSONEncoder)


def replace_placeholders(template_path, context, output_path):
    # Load the template document
    if not os.path.exists(template_path):
        print(f"Error: The file at {template_path} does not exist.")
        return None
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Error loading template: {e}")
        return None
    
    # Iterate over each paragraph in the document
    for paragraph in doc.paragraphs:
        for placeholder, value in context.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    # Iterate over each table in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in context.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
    
    # Save the modified document to the specified output path
    doc.save(output_path)
    
    return output_path

def generate_report(report_df, template_path, file_path, test_engineer):
    today = datetime.today()
    today_string = today.strftime("%d/%m/%Y")
    today=today.strftime("%d%m%Y")
    # Iterate over the DataFrame rows
    for index, row in report_df.iterrows():
        
        report_no = f"{row['PROJECT ID']}_{today}_v0"
        
        # Format the period without time using pd.Timestamp
        row["Test Start Date"] = pd.Timestamp(row["START DATE"]).strftime("%d/%m/%Y")
        row["Test End Date"] = pd.Timestamp(row["END DATE"]).strftime("%d/%m/%Y")
        
        period = f"{row['START DATE']} - {row['END DATE']}"
        if test_engineer:
            parts = test_engineer.lower().split(" ")
            
            email = parts[0] + "." + parts[1] + "@gmail.com"
            context = {
   
                "{PROJECT ID}": row["PROJECT ID"],
                "{REPORT_NO}": report_no,
                "{PERIOD}": period,
                "{MAIL}": email,

            }
            file_name = report_no
            # Generate a unique output file path
            output_file_path = os.path.join(os.path.dirname(file_path), f"{file_name}.pdf")
        

            # Replace placeholders in the template and save to the new file
            replace_placeholders(template_path, context,file_path)
        else:
            raise Exception("Who is the engineer?")
            

def load_template(test_value):
        # Get the absolute path of the current script
    current_dir = os.path.dirname(os.path.abspath(__file__))
    test_value = test_value.strip()
    
    # Construct the full path to the template directory
    template_dir = os.path.join(current_dir, "TEMPLATES")
    template_path = os.path.join(template_dir, f"{test_value}.docx")
    
    if os.path.exists(template_path):
        return template_path
    else:
        st.error(f"Template for {test_value} not found.")
        return None
