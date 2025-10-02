import os, re, shutil, json, sqlite3, plotly
import pandas as pd
import streamlit as st
from io import BytesIO
from datetime import datetime
from sqlalchemy import create_engine
from typing import List, Dict, Literal
import plotly.figure_factory as ff
import plotly.graph_objects as go
import plotly.express as px
import matplotlib.pyplot as plt
from docx import Document
 
from database import *   

st.set_page_config(
    page_title="PRODUCT ENGINEERING - VALIDATION",
    layout="wide",
)

# --- Validation Tracker Class ---
class ValidationTracker:
    def __init__(self):
        database()
        self.query = "SELECT * FROM ProjectTracker"
        self.data = self.load_data()
        self.column_config = self.get_column_config()

    def load_data(self) -> pd.DataFrame:
        data = get_data_from_db(self.query)

        # Enforce data types
        if 'Day' in data.columns:
            data['Day'] = pd.to_datetime(data['Day'], errors='coerce')
        for col in ['Datasheet', 'Function', 'EMC']:
            if col in data.columns:
                data[col] = data[col].astype(bool)
        if 'Homologated' not in data.columns:
            data['Homologated'] = ""
        data['Progress'] = data.apply(
            lambda row: 0 if row["Homologated"] in ["Passed", "Failed"]
            else (pd.Timestamp.now() - row["Day"]).days if pd.notnull(row["Day"]) else 0,
            axis=1
        )
        data['Progress'] = data['Progress'].astype(int)
        return data

    def get_column_config(self) -> Dict[str, st.column_config.Column]:
        return {
            "Datasheet": st.column_config.CheckboxColumn("Datasheet", default=False),
            "Function": st.column_config.CheckboxColumn("Function", default=False),
            "EMC": st.column_config.CheckboxColumn("EMC", default=False),
            "Progress": st.column_config.ProgressColumn(
                "Progress", min_value=0, max_value=40, format="%.0f days"
            ),
            "Homologated": st.column_config.TextColumn(
                "Homologated", help="Status of validation", validate=".*", required=False
            )
        }

    def display_editor(self) -> pd.DataFrame:
        return st.data_editor(
            self.data,
            hide_index=True,
            num_rows="dynamic",
            column_config=self.column_config
        )

    def save_changes(self, edited_data: pd.DataFrame):
        engine = create_engine('sqlite:///project_tracker.db')
        edited_data.to_sql('ProjectTracker', con=engine, if_exists='replace', index=False)
        st.success("Changes saved successfully.")

    def download_backup(self, edited_data: pd.DataFrame):
        backup = BytesIO()
        with pd.ExcelWriter(backup) as writer:
            edited_data.to_excel(writer, index=False)
        today = datetime.today().strftime("%d%m%Y")
        st.download_button(
            label="🗂️ Download Backup",
            data=backup.getvalue(),
            file_name=f"Backup_Project_Tracker_{today}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    def display_charts(self):
        df = self.data

        # Count Checked and Unchecked for each category
        datasheet_counts = df['Datasheet'].value_counts().rename({True: 'Checked', False: 'Unchecked'})
        function_counts = df['Function'].value_counts().rename({True: 'Checked', False: 'Unchecked'})
        emc_counts = df['EMC'].value_counts().rename({True: 'Checked', False: 'Unchecked'})

        # Create grouped bar chart
        fig = go.Figure()

        # Validation metrics (left axis)
        fig.add_trace(go.Bar(name='Datasheet Checked', x=['Datasheet'], y=[datasheet_counts.get('Checked', 0)], marker_color='lightgreen'))
        fig.add_trace(go.Bar(name='Datasheet Unchecked', x=['Datasheet'], y=[datasheet_counts.get('Unchecked', 0)], marker_color='salmon'))

        fig.add_trace(go.Bar(name='Function Checked', x=['Function'], y=[function_counts.get('Checked', 0)], marker_color='lightgreen'))
        fig.add_trace(go.Bar(name='Function Unchecked', x=['Function'], y=[function_counts.get('Unchecked', 0)], marker_color='salmon'))

        fig.add_trace(go.Bar(name='EMC Checked', x=['EMC'], y=[emc_counts.get('Checked', 0)], marker_color='lightgreen'))
        #fig.add_trace(go.Bar(name='EMC Unchecked', x=['EMC'], y=[emc_counts.get('Unchecked', 0)], marker_color='tomato'))
        
            # Layout
        fig.update_layout(
                title="Validation Summary",
                barmode='group',
                xaxis_title="Category",
                yaxis=dict(title='Validation Counts'),
                width=800,
                height=600
            )

            # Display in Streamlit
        st.plotly_chart(fig, use_container_width=True)


    def replace_placeholders(template_path, context, output_path):
        # Load the template document
        if not os.path.exists(template_path):
            print(f"Error: The file at {template_path} does not exist.")
            return None
        try:
            doc = Document(template_path)
        except Exception as e:
            print(f"Error loading template: {e}")
            return None
        
        # Iterate over each paragraph in the document
        for paragraph in doc.paragraphs:
            for placeholder, value in context.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Iterate over each table in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in context.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
        
        # Save the modified document to the specified output path
        doc.save(output_path)
        
        return output_path

    def generate_report(report_df, template_path, file_path, test_engineer):
        today = datetime.today()
        today_string = today.strftime("%d/%m/%Y")
        today=today.strftime("%d%m%Y")
        # Iterate over the DataFrame rows
        for index, row in report_df.iterrows():
            
            report_no = f"{row['PROJECT ID']}_{today}_v0"
            
            # Format the period without time using pd.Timestamp
            row["Test Start Date"] = pd.Timestamp(row["START DATE"]).strftime("%d/%m/%Y")
            row["Test End Date"] = pd.Timestamp(row["END DATE"]).strftime("%d/%m/%Y")
            
            period = f"{row['START DATE']} - {row['END DATE']}"
            if test_engineer:
                parts = test_engineer.lower().split(" ")
                
                email = parts[0] + "." + parts[1] + "@gmail.com"
                context = {
    
                    "{PROJECT ID}": row["PROJECT ID"],
                    "{REPORT_NO}": report_no,
                    "{PERIOD}": period,
                    "{MAIL}": email,

                }
                file_name = report_no
                # Generate a unique output file path
                output_file_path = os.path.join(os.path.dirname(file_path), f"{file_name}.pdf")
            

                # Replace placeholders in the template and save to the new file
                replace_placeholders(template_path, context,file_path)
            else:
                raise Exception("Who is the engineer?")
                

    def load_template(test_value):
            # Get the absolute path of the current script
        current_dir = os.path.dirname(os.path.abspath(__file__))
        test_value = test_value.strip()
        
        # Construct the full path to the template directory
        template_dir = os.path.join(current_dir, "TEMPLATES")
        template_path = os.path.join(template_dir, f"{test_value}.docx")
        
        if os.path.exists(template_path):
            return template_path
        else:
            st.error(f"Template for {test_value} not found.")
            return None


# --- To-Do Manager Class ---
class TodoManager:
    TODO_FILE = "todo_list.json"
    PRIORITY_LEVELS = ["High", "Medium", "Low"]

    def load_todo(self) -> List[Dict[str, str]]:
        if os.path.exists(self.TODO_FILE):
            with open(self.TODO_FILE, "r") as f:
                todos = json.load(f)
                valid_todos = []
                for item in todos:
                    task = str(item.get("task", "")).strip()
                    priority = item.get("priority", "Medium")
                    due_date = item.get("due_date", "")
                    if priority not in self.PRIORITY_LEVELS:
                        priority = "Medium"
                    if task:
                        valid_todos.append({
                            "task": task,
                            "priority": priority,
                            "due_date": due_date
                        })
                return valid_todos
        return []

    def save_todo(self, todos: List[Dict[str, str]]):
        with open(self.TODO_FILE, "w") as f:
            json.dump(todos, f, indent=2)

    def display_calendar(self):
        todos = self.load_todo()
        if not todos:
            st.info("No tasks scheduled.")
            return

        df = pd.DataFrame(todos)
        df["due_date"] = pd.to_datetime(df["due_date"], errors="coerce")

        # Sort by priority and due date
        priority_order = {"High": 1, "Medium": 2, "Low": 3}
        df["priority_rank"] = df["priority"].map(priority_order)
        df = df.sort_values(["priority_rank", "due_date"])

        priority_colors = {
            "High": "#902018",   # Red
            "Medium": "#93871c", # Yellow
            "Low": "#2b6a2d"     # Green
        }

        # Create 3 columns for Low, Medium, High
        col_low, col_medium, col_high = st.columns(3)

        for priority, col in zip(["Low", "Medium", "High"], [col_low, col_medium, col_high]):
            with col:
                st.markdown(f"### {priority}")
                priority_tasks = df[df["priority"] == priority]

                if priority_tasks.empty:
                    st.markdown("No tasks.")
                else:
                    for i, row in priority_tasks.iterrows():
                        due_date = row["due_date"]
                        date_str = due_date.strftime('%d %b %Y') if pd.notnull(due_date) else "No due date"
                        task_key = f"{row['task']}_{i}"

                        st.markdown(f"""
                        <div style="background-color:{priority_colors[priority]}; padding:4px; border-radius:4px; margin-bottom:4px; font-size:12px; color:white;">
                            📝 <strong>{row['task']}</strong><br>
                            📆 <em>{date_str}</em>
                        </div>
                        """, unsafe_allow_html=True)

                        with st.expander("❌ Cancel", expanded=False):
                            confirm_key = f"confirm_{task_key}"
                            if st.checkbox(f"Confirm cancel '{row['task']}'", key=confirm_key):
                                todos.remove({
                                    "task": row["task"],
                                    "priority": row["priority"],
                                    "due_date": row["due_date"].strftime("%Y-%m-%d") if pd.notnull(row["due_date"]) else ""
                                })
                                self.save_todo(todos)
                                st.success(f"Task '{row['task']}' cancelled.")
                                
    
    def add_task(self):
            new_task = st.text_input("Enter a new task")
            priority = st.selectbox("Select priority", self.PRIORITY_LEVELS)
            due_date = st.date_input("Select due date", value=datetime.today())
            if st.button("Save Task"):
                if new_task.strip():
                    todos = self.load_todo()
                    todos.append({
                        "task": new_task.strip(),
                        "priority": priority,
                        "due_date": due_date.strftime("%Y-%m-%d")
                    })
                    self.save_todo(todos)
                    st.success("Task saved successfully!")
                else:
                    st.warning("Please enter a valid task.")

    


# --- Main App ---
def project_tracker():
    
    with st.sidebar:
        st.markdown("📊 Validation tracker by Vincentiu")
        uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")
        debug_mode = st.checkbox("🔧 Developer Mode", value=False)
        if uploaded_file:
            fill_database(uploaded_file)
            st.success("Database has been populated successfully.")

    tab1, tab2 = st.tabs(["📊 Validation request", "📥Todo!"])

    with tab1:
        tracker = ValidationTracker()
        but1, but2, but3, but4 = st.columns(4, gap="small")
        st.text("MOS, Diodes and all resonant components need EMC & Functionality test")
        edited_data = tracker.display_editor()
        tracker.display_charts()
        with but1:
            if st.button("💾 Save Changes"):
                tracker.save_changes(edited_data)

        with but2:
            tracker.download_backup(edited_data)

        
        if debug_mode:
                with but3:
                    if st.button("📥 High volume reports"):
                        tracker.generate_reports(edited_data)

                with but4:
                    if st.button("💼 Single report"):
                        st.info("Single report generation logic goes here.")


    
    with tab2:
        todo = TodoManager()
        with st.expander("Add task"):
            todo.add_task()
        todo.display_calendar()
        


project_tracker()